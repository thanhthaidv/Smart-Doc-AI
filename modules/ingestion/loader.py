import os

import pandas as pd
from docx import Document as DocxDocument
import pdfplumber
from langchain_core.documents import Document
from modules.ingestion.ocr import (
    detect_pdf_has_images,
    detect_pdf_has_text,
    ocr_pdf_pages_to_text,
)

def load_pdf(file_path: str, use_ocr_if_needed: bool = True):
    has_images = detect_pdf_has_images(file_path, max_pages=100) # 
    has_text = detect_pdf_has_text(file_path) # 
    documents = []

    # Sử dụng pdfplumber mở trực tiếp để kiểm soát từng trang
    with pdfplumber.open(file_path) as pdf:
        # Nếu cần OCR, thực hiện một lần để lấy danh sách text từng trang
        ocr_pages = []
        if use_ocr_if_needed and (not has_text or has_images):
            ocr_pages = ocr_pdf_pages_to_text(file_path) # 

        for idx, page in enumerate(pdf.pages):
            page_number = idx + 1
            # Trích xuất text layer (nếu có)
            page_text = (page.extract_text() or "").strip() # 
            
            # Lấy text OCR tương ứng với trang này
            current_ocr_text = ocr_pages[idx] if idx < len(ocr_pages) else "" # 
            
            # Hợp nhất nội dung: Ưu tiên Text layer, bổ sung OCR nếu có ảnh
            final_content = page_text
            is_ocr_used = False
            
            if use_ocr_if_needed and current_ocr_text.strip():
                if not final_content: # Trang trắng/thuần ảnh như trang 5 của bạn
                    final_content = current_ocr_text
                else: # Trang hybrid (vừa có chữ vừa có ảnh)
                    final_content = f"{final_content}\n\n[OCR Content]:\n{current_ocr_text}"
                is_ocr_used = True

            if not final_content.strip():
                continue

            # Tạo Document với metadata chính xác tuyệt đối theo vòng lặp
            documents.append(
                Document(
                    page_content=final_content,
                    metadata={
                        "source": file_path,
                        "page": page_number,
                        "chunk_id": page_number,
                        "ocr": is_ocr_used,
                        "has_images": has_images
                    }
                )
            )
            
    return documents


def load_docx(file_path: str):
    """Tải DOCX và trả về nội dung đã hợp nhất."""
    doc = DocxDocument(file_path)
    parts = [p.text for p in doc.paragraphs if p.text.strip()]

    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                parts.append("\t".join(cells))

    text = "\n".join(parts)
    return [
        Document(
            page_content=text,
            metadata={"source": file_path, "type": "docx", "chunk_id": 1},
        )
    ]


def load_file(file_path: str, use_ocr_if_needed: bool = True):
    """Điều phối để tải nhiều định dạng file."""
    ext = os.path.splitext(file_path)[1].lower()
    if ext == ".pdf":
        return load_pdf(file_path, use_ocr_if_needed=use_ocr_if_needed)
    if ext == ".docx":
        return load_docx(file_path)

    raise ValueError(f"Unsupported file type: {ext}")

